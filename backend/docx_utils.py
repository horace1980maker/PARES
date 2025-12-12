"""
Utility module to convert Markdown text to Word (.docx) documents
"""
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_docx_from_markdown(markdown_text: str) -> BytesIO:
    """
    Convert markdown text to a Word document and return as BytesIO buffer
    
    Args:
        markdown_text: Markdown-formatted text to convert
        
    Returns:
        BytesIO buffer containing the DOCX file
    """
    doc = Document()
    
    # Split by lines
    lines = markdown_text.split('\n')
    
    in_code_block = False
    code_block_lines = []
    in_list = False
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Quote'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(40, 40, 40)
                code_block_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # Handle headers
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            in_list = False
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            in_list = False
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            in_list = False
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            in_list = False
        
        # Handle bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = process_inline_markdown(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            in_list = True
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = process_inline_markdown(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            in_list = True
        
        # Handle regular paragraphs
        elif line.strip():
            text = process_inline_markdown(line.strip())
            p = doc.add_paragraph()
            add_formatted_text(p, text)
            in_list = False
        else:
            # Empty line
            if in_list:
                in_list = False
    
    # Save to memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def process_inline_markdown(text: str) -> str:
    """
    Process inline markdown (bold, italic, code) and return text with markers
    This is a simplified version - we'll handle formatting when adding to paragraph
    """
    return text


def add_formatted_text(paragraph, text: str):
    """
    Add text to paragraph with inline formatting for bold, italic, code
    """
    # Pattern to match **bold**, *italic*, and `code`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(200, 0, 0)
        else:
            # Regular text
            paragraph.add_run(part)
