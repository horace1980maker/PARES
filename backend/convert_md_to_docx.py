"""
Script to convert markdown to Word (.docx) document
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(md_file, output_file):
    """Convert markdown to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Split by lines
    lines = content.split('\n')
    
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Quote'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_block_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle tables
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        else:
            # Process accumulated table
            if in_table:
                process_table(doc, table_lines)
                in_table = False
                table_lines = []
            
            # Handle regular paragraphs
            if line.strip():
                # Check for bold
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                
                p = doc.add_paragraph(line)
                
                # Handle bullet points
                if line.strip().startswith('- '):
                    p.style = 'List Bullet'
                elif line.strip().startswith('* '):
                    p.style = 'List Bullet'
                elif re.match(r'^\d+\.', line.strip()):
                    p.style = 'List Number'
    
    # Save document
    doc.save(output_file)
    print(f"Document saved to: {output_file}")

def process_table(doc, table_lines):
    """Process table lines and add to document"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Skip separator line
    if len(table_lines) < 3:
        return
    
    # Parse rows
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    
    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Add rows
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i + 1].cells[j].text = cell

if __name__ == "__main__":
    md_file = r"C:\Users\narva\.gemini\antigravity\brain\4ec852d4-662f-4133-b498-f3ed12041555\especificaciones_tecnicas_rag.md"
    output_file = r"C:\Users\narva\.gemini\antigravity\brain\4ec852d4-662f-4133-b498-f3ed12041555\especificaciones_tecnicas_rag.docx"
    
    try:
        parse_markdown_to_docx(md_file, output_file)
    except ImportError:
        print("python-docx is not installed. Installing...")
        import subprocess
        subprocess.run(["pip", "install", "python-docx"])
        parse_markdown_to_docx(md_file, output_file)
